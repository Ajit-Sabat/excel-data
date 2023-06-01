import docx

doc=docx.Document('message.docx')

para=doc.add_paragraph()
run=para.add_run("13880")
run.font.strike=True

print(run)